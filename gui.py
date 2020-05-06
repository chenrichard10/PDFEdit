import pdf_edit as edit
import tkinter as tk
from tkinter import filedialog
import os
from PyPDF2 import PdfFileReader, PdfFileWriter
from docx import Document
from docx.shared import Inches
import sys
import comtypes.client
# os.startfile(opens the specified )
class Application(tk.Tk):
    def __init__(self):
        tk.Tk.__init__(self)
        self.resizable(0, 0)
        self.geometry("400x400")
        self.title("PDFEdit")
        self.iconbitmap('C:/Users/chenp/Documents/GitHub/PDFEdit/res/pdf.ico')
        self.create_widgets()

    # This function opens 3 separate file dialogs to find the combined PDF and location 
    # to be saved
    # Merged file is opened afterwards 
    def merge(self):
        pdf1 = tk.filedialog.askopenfilename(initialdir = os.getcwd(), title = "Select first PDF")
        pdf2 = tk.filedialog.askopenfilename(initialdir = os.getcwd(), title = "Select second PDF")
        location = tk.filedialog.askdirectory(initialdir = os.getcwd(), title = "Select a folder to be stored")
        combined = edit.merge_pdfs('merged.pdf', location, pdf1, pdf2)
        os.startfile(combined)

    def pdfToDocx(self):
        document = Document()
        document.add_heading('Document Title', 0)
        text=""
        pdf_file = tk.filedialog.askopenfilename(initialdir = os.getcwd(), title = "Select PDF")
        pdf_file
        read_pdf = PdfFileReader(pdf_file)
        c = read_pdf.numPages
        print('gunga')
        for i in range(c):
            page = read_pdf.getPage(i)
            text+=(page.extractText())
        
        document.add_paragraph(text)
        document.add_page_break()
        document.save(read_pdf.documentInfo.title + '.docx')

    def docxtoPDF(self):
        wdFormatPDF = 17
        word = comtypes.client.CreateObject('Word.Application')
        doc = tk.filedialog.askopenfilename(initialdir = os.getcwd(), title = "Select first PDF")
        doc.SaveAs(doc.Title + '.pdf', FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit() 

    
    def create_widgets(self):
        # Creating the buttons 
        self.merge_btn = tk.Button(self, text = "Merge PDF", padx = 50, pady = 50, command = self.merge)
        self.convert_btn = tk.Button(self, text = "Convert to .docx", padx = 45, pady = 50, command=self.docxtoPDF)
        self.edit_btn = tk.Button(self, text = "Edit PDF", padx = 56, pady = 50)
        self.extract_btn = tk.Button(self, text = "Convert to .pdf", padx = 64, pady = 50)
        # Positioning the buttons 
        self.merge_btn.grid(row = 2, column = 2,  padx = 10, pady = 20)
        self.convert_btn.grid(row = 2, column = 4,  padx = 10, pady = 20)
        self.edit_btn.grid(row = 4, column = 2,  padx = 10, pady = 20)
        self.extract_btn.grid(row = 4, column = 4,  padx = 10, pady = 20)

if __name__ == '__main__':
    app = Application()
    app.mainloop()

