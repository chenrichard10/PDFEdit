from tkinter import *
from PyPDF2 import PdfFileReader, PdfFileWriter
from docx import Document
from docx.shared import Inches
import sys
import os
import comtypes.client


document = Document()
document.add_heading('Document Title', 0)
window = Tk() 
window.geometry('400x150')  
window.title('PDF Editor')
# greeting = tk.Label(text = "Richard is a gunga")
# greeting.pack()
# def test() :
#     print('gunga')



def pdfToDocx(path):
    text=""
    pdf_file = open(path, 'rb')
    read_pdf = PdfFileReader(pdf_file)
    c = read_pdf.numPages
    print('gunga')
    for i in range(c):
         page = read_pdf.getPage(i)
         text+=(page.extractText())
    
    document.add_paragraph(text)
    document.add_page_break()
    document.save(path + '.docx')

def docxtoPDF(path):
    wdFormatPDF = 17
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(path)
    doc.SaveAs(path + '.pdf', FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit() 


button = Button(window, 
    text="Clicked GGez",
    bg='black',
    fg='blue',
    width=25,
    height=5,
    # command=lambda: pdfToDocx('Felix Peer Evaluations Feedback (1).pdf')
    command = lambda: docxtoPDF('Tset.docx')
)
button.pack()
window.mainloop()


